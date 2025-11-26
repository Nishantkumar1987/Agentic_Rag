import os
import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from docx import Document

# --------------------------------------------
# FINAL LOCKED CANONICAL TOPICS (FOR ACCOUNTS)
# --------------------------------------------
CANONICAL_TOPICS = [
    "Product Overview",
    "Features & Benefits",
    "Eligibility",
    "KYC / Documentation",
    "Fees & Charges",
    "Interest / Pricing",
    "Eligibility Addendum",
    "MITC",
    "Transaction & Usage Rules",
    "Limits",
    "Instruments & Tools",
    "Statements & Communication",
    "Dormancy / Inoperative / Surrender",
    "Closure",
    "Complaints & Grievances",
    "Legal T&C"
]

# --------------------------------------------
# STRONGER NORMALIZATION RULES
# --------------------------------------------
def normalize_section_name(raw):
    if not raw:
        return None

    t = raw.lower().strip()

    # Strong keyword-based canonicalization
    if "feature" in t:
        return "Features & Benefits"

    if "eligibility addendum" in t:
        return "Eligibility Addendum"

    if "eligibility" in t:
        return "Eligibility"

    if "kyc" in t or "documentation" in t:
        return "KYC / Documentation"

    if "interest" in t or "pricing" in t or "rate" in t:
        return "Interest / Pricing"

    if "service charge" in t or "fee" in t or "charges" in t:
        return "Fees & Charges"

    if "important terms" in t or "mitc" in t:
        return "MITC"

    if "withdraw" in t or "transaction" in t or "usage" in t:
        return "Transaction & Usage Rules"

    if "limit" in t:
        return "Limits"

    if "cheque" in t or "debit card" in t or "instruments" in t:
        return "Instruments & Tools"

    if "statement" in t or "communication" in t or "sms" in t:
        return "Statements & Communication"

    if "dormant" in t or "inoperative" in t or "surrender" in t:
        return "Dormancy / Inoperative / Surrender"

    if "closure" in t or "close account" in t:
        return "Closure"

    if "complaint" in t or "griev" in t or "feedback" in t:
        return "Complaints & Grievances"

    if "legal" in t or "term" in t or "condition" in t:
        return "Legal T&C"

    # fallback: product overview-like
    if "overview" in t or "about" in t:
        return "Product Overview"

    return raw.strip().title()  # fallback

# --------------------------------------------
# IMPROVED HEADING DETECTION
# --------------------------------------------
def detect_heading(text):
    if not text:
        return False

    t = text.strip()

    # Strong signals
    if t.isupper():                      # Entire heading uppercase (common)
        return True

    if t.endswith(":"):                  # Ends with colon
        return True

    # Short bold / capital headings
    if len(t.split()) <= 7 and t[0].isupper():
        return True

    # If it matches part of canonical topics
    for c in CANONICAL_TOPICS:
        if c.split()[0].lower() in t.lower():
            return True

    # Look for common patterns in account docs
    patterns = ["features", "eligibility", "interest", "service", "charges",
                "dormant", "statement", "kyc", "mitc"]
    if any(p in t.lower() for p in patterns):
        return True

    return False

# --------------------------------------------
# TABLE PARSER
# --------------------------------------------
def convert_table(table):
    rows = []
    header = [c.text.strip() for c in table.rows[0].cells]
    if all(h == "" for h in header):
        header = [f"col_{i+1}" for i in range(len(header))]
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        rows.append({header[i]: cells[i] for i in range(len(header))})
    return rows

# --------------------------------------------
# MAIN PARSER FUNCTION
# --------------------------------------------
def parse_account_docx(doc_path, output_dir="parsed"):
    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    sections = []
    current = None

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue

        if detect_heading(text):
            # close previous section
            if current:
                sections.append(current)

            normalized = normalize_section_name(text)

            current = {
                "section_id": str(uuid.uuid4()),
                "title": normalized,
                "type": normalized,
                "text": "",
                "tables": [],
                "status": "present"
            }
        else:
            if current:
                current["text"] += "\n" + text

    if current:
        sections.append(current)

    # --------------------------------------------
    # TABLE ATTACHMENT (attach to Fees or MITC preferably)
    # --------------------------------------------
    for tbl in doc.tables:
        tjson = convert_table(tbl)

        target = None
        for sec in sections:
            if sec["type"] in ["Fees & Charges", "MITC"]:
                target = sec
                break

        if not target:
            target = sections[0]

        target["tables"].append({
            "table_id": str(uuid.uuid4()),
            "json": tjson
        })

    # --------------------------------------------
    # ADD MISSING CANONICAL SECTIONS
    # --------------------------------------------
    present = {s["type"] for s in sections}

    for topic in CANONICAL_TOPICS:
        if topic not in present:
            sections.append({
                "section_id": str(uuid.uuid4()),
                "title": topic,
                "type": topic,
                "text": "",
                "tables": [],
                "status": "missing"
            })

    # --------------------------------------------
    # BUILD JSON FILE
    # --------------------------------------------
    file_name = Path(doc_path).name
    product_name = Path(doc_path).stem
    product_id = re.sub(r"[^a-z0-9]+", "_", product_name.lower()).strip("_")

    final_json = {
        "product_id": product_id,
        "product_name": product_name,
        "product_line": "Account",
        "documents": [
            {
                "file_name": file_name,
                "source_type": "docx",
                "parsed_at": datetime.now().isoformat(),
                "sections": sections
            }
        ]
    }

    Path(output_dir).mkdir(exist_ok=True)
    out_path = Path(output_dir) / f"{product_id}.json"

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(final_json, f, indent=2, ensure_ascii=False)

    print("Parsed:", out_path)
    return out_path

# --------------------------------------------
# BATCH PARSING
# --------------------------------------------
def parse_all_accounts_in_folder(folder_path, output_dir="parsed"):
    folder = Path(folder_path)
    for file in folder.iterdir():
        if file.suffix.lower() == ".docx":
            print("Processing:", file.name)
            parse_account_docx(str(file), output_dir=output_dir)

    print("\n✓ All files parsed.\nJSON saved in:", output_dir)
